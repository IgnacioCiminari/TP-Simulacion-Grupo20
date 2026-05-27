import sys
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

path = r'c:\Users\Usuario\Documents\Facultad - UTN - Ingeniería en Sistemas de Información\Cuarto Año\Simulacion\U5 - Modelos de Simulación Dinámicos\DocumentoWord\Analisis de simulación.docx'

try:
    doc = docx.Document(path)
except Exception as e:
    print(f"Error opening document: {e}")
    sys.exit(1)

# Find 'Análisis' paragraph
analysis_idx = -1
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().lower() == "análisis":
        analysis_idx = i
        break

if analysis_idx != -1:
    # Delete paragraphs after
    for p in doc.paragraphs[analysis_idx+1:]:
        p._element.getparent().remove(p._element)
    
    # Delete tables
    for t in doc.tables:
        t._element.getparent().remove(t._element)

def add_bullet(text, indent=0.25):
    p = doc.add_paragraph(f"• {text}")
    p.paragraph_format.left_indent = Inches(indent)
    return p

def add_num(text, indent=0.25):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(indent)
    return p

doc.add_heading('1. Entidades y Estados', level=2)

doc.add_heading('1.1. Cliente (Vehículo)', level=3)
doc.add_paragraph('Representa la entidad temporal que fluye a través del sistema.')
doc.add_paragraph('Tipos:')
add_bullet('Camioneta (Prioridad Alta)')
add_bullet('Auto (Prioridad Baja)')

doc.add_paragraph('Estados Posibles:')
add_bullet('Esperando: En la cola de entrada al predio.')
add_bullet('En Frenos: Siendo atendido en la estación de Frenos.')
add_bullet('Bloqueado en Frenos: Terminó su revisión de Frenos pero no puede avanzar porque Luces está ocupada.')
add_bullet('En Luces: Siendo atendido en la estación de Luces.')
add_bullet('Retirado: Revisión finalizada, fuera del sistema.')

doc.add_heading('1.2. Servidores (Infraestructura)', level=3)
doc.add_paragraph('La planta se compone de Líneas de Inspección (L_i), cada una conformada por dos estaciones secuenciales:')

doc.add_heading('Estación de Frenos (Frenos L_i)', level=4)
doc.add_paragraph('Estados Posibles:')
add_bullet('Libre: Sin vehículo, lista para atender.')
add_bullet('Ocupado: Atendiendo a un vehículo.')
add_bullet('Bloqueado: El vehículo terminó, pero no puede avanzar. La estación no puede tomar clientes nuevos.')

doc.add_heading('Estación de Luces y Emisiones (Luces L_i)', level=4)
doc.add_paragraph('Estados Posibles:')
add_bullet('Libre: Sin vehículo, lista para atender.')
add_bullet('Ocupado: Atendiendo a un vehículo.')

doc.add_heading('1.3. Cola de Espera (Entry Queue)', level=3)
add_bullet('Es una Cola de Prioridad Única.')
add_bullet('Las Camionetas se ubican por delante de los Autos. A igual tipo de vehículo, se respeta el orden de llegada (FIFO).')

doc.add_heading('2. Variables Aleatorias', level=2)
doc.add_paragraph('El sistema presenta variabilidad en las llegadas y en los tiempos de servicio. Se utiliza el Método de la Transformada Inversa:')

# Table
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Variable'
hdr_cells[1].text = 'Distribución'
hdr_cells[2].text = 'Fórmula Generadora'

data = [
    ('Tiempo entre Llegada Autos', 'Exponencial (μ=15)', 'T = -15 × ln(1 - RND)'),
    ('Tiempo entre Llegada Camionetas', 'Exponencial (μ=30)', 'T = -30 × ln(1 - RND)'),
    ('Tiempo de Revisión Frenos', 'Uniforme (A=4, B=7)', 'T = 4 + RND × 3'),
    ('Tiempo de Revisión Luces', 'Uniforme (A=6, B=10)', 'T = 6 + RND × 4'),
]
for var, dist, form in data:
    row_cells = table.add_row().cells
    row_cells[0].text = var
    row_cells[1].text = dist
    row_cells[2].text = form

doc.add_heading('3. Eventos Discretos y Lógica de Transición', level=2)

doc.add_heading('Llegada Auto / Llegada Camioneta', level=3)
add_bullet('Condición: Ocurren continuamente hasta el Cierre de Puertas.', indent=0)
add_bullet('Lógica:', indent=0)
add_num('1. Ingresa al predio. Genera de inmediato su próxima llegada correspondiente.', indent=0.25)
add_num('2. Evalúa si existe alguna estación de Frenos Libre en cualquier línea (priorizando la Línea 1 por defecto).', indent=0.25)
add_num('3. Si hay Frenos libres: El vehículo entra a la estación, pasa a estado En Frenos, la estación pasa a Ocupado y se programa el evento Fin Revisión Frenos.', indent=0.25)
add_num('4. Si no hay Frenos libres: El vehículo va a la cola de espera, respetando su nivel de prioridad.', indent=0.25)

doc.add_heading('Fin Revisión Frenos L_i', level=3)
add_bullet('Condición: Se dispara cuando finaliza el tiempo generado por la distribución uniforme de Frenos para la línea i.', indent=0)
add_bullet('Lógica:', indent=0)
add_num('1. El vehículo termina la primera etapa.', indent=0.25)
add_num('2. Evalúa la estación de Luces de su misma línea.', indent=0.25)
add_num('3. Si Luces está Libre: Pasa a En Luces, Luces se marca Ocupado, Frenos queda Libre. Se programa el evento Fin Revisión Luces. Se revisa la cola de entrada: si hay vehículos, ingresa el de mayor prioridad a Frenos.', indent=0.25)
add_num('4. Si Luces está Ocupado: El vehículo queda Bloqueado en Frenos, y la estación de Frenos pasa a Bloqueado.', indent=0.25)

doc.add_heading('Fin Revisión Luces L_i', level=3)
add_bullet('Condición: Finaliza el tiempo de la segunda etapa.', indent=0)
add_bullet('Lógica:', indent=0)
add_num('1. El vehículo sale del sistema (Retirado). Se registran sus tiempos en las estadísticas globales.', indent=0.25)
add_num('2. La estación de Luces pasa a Libre.', indent=0.25)
add_num('3. Desbloqueo: Evalúa la estación de Frenos de su misma línea. Si estaba Bloqueado, el vehículo bloqueado avanza automáticamente a Luces, liberando Frenos. Si Frenos se liberó, toma al siguiente vehículo de la cola de entrada (si hubiera).', indent=0.25)

doc.add_heading('Cierre de Puertas', level=3)
add_bullet('Condición: Ocurre a los 960 minutos (16:00 hs).', indent=0)
add_bullet('Lógica: Cancela las futuras llegadas. Ningún nuevo cliente entra al predio, pero la simulación continúa procesando la cola y los vehículos en curso.', indent=0)

doc.add_heading('Fin de Simulación (Condición de Cierre)', level=3)
doc.add_paragraph('No es un evento programado con tiempo estático, sino una consecuencia natural (Condición de Corte) del DES. Ocurre cuando la lista de eventos futuros (FEL) se vacía, garantizando que el último vehículo fue despachado.')

p = doc.add_paragraph()
r = p.add_run('Nota: Este análisis está alineado con la implementación final en código, optimizando algunos conceptos (como la parametrización de eventos por línea y la unificación de colas por prioridad) respecto al diseño original.')
r.italic = True
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.save(path)
print("Done!")
